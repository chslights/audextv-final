"""
audit_ingestion_v04.2/audit_ingestion/extractor.py
Page-aware extraction engine with fast/deep lane split.

Fast Review (extract_fast):
  - pdfplumber text only (no table scan)
  - shared PyPDF2 reader for weak pages
  - no OCR, no vision
  - escalate globally if total extraction is weak

Deep Extraction (extract_deep):
  - pdfplumber text + targeted table scan
  - shared PyPDF2 reader
  - OCR on weak pages (top 6 max)
  - vision on critical pages (top 2 max)

Shared PDF handles — never reopen per page.
"""
from __future__ import annotations
import hashlib
import io
import logging
from pathlib import Path
from typing import Optional
from .models import ParsedDocument, ParsedPage, ParsedTable,\
    VISION_PAGE_NOT_ATTEMPTED, VISION_PAGE_BASE_EXTRACTED,\
    VISION_PAGE_REQUESTED, VISION_PAGE_COMPLETE,\
    VISION_PAGE_PARTIAL, VISION_PAGE_FAILED, VISION_PAGE_REVIEW_NEEDED,\
    MERGE_KEEP_NATIVE, MERGE_REPLACE, MERGE_APPEND,\
    TRIGGER_WEAK_NATIVE, TRIGGER_MISSING_FIELDS, TRIGGER_HANDWRITING_SUSPECTED,\
    TRIGGER_LOW_NUMERIC, TRIGGER_IMAGE_FILE,\
    ERR_VISION_NO_IMPROVEMENT, ERR_VISION_OUTPUT_EMPTY

logger = logging.getLogger(__name__)

# ── Limits ────────────────────────────────────────────────────────────────────
MAX_PDF_PAGES_FAST        = 40
MAX_PDF_PAGES_TABLE_SCAN  = 12
MAX_OCR_PAGES             = 6
MAX_VISION_PAGES          = 2
MAX_RELEVANT_PAGES        = 10
MAX_CONTEXT_CHARS         = 20000

# ── Thresholds ────────────────────────────────────────────────────────────────
MIN_CHARS_ACCEPTABLE  = 350
MIN_CHARS_WEAK        = 150
MIN_CHARS_CRITICAL    = 60


# ── Vision mode constants ─────────────────────────────────────────────────────
VISION_AUTO            = "auto"            # existing behavior — vision only on critical weak pages
VISION_STANDARD_ONLY   = "standard_only"   # no vision at all
VISION_FORCE_ALL       = "force_all"       # vision on every page
VISION_FORCE_SELECTED  = "force_selected"  # vision on user-specified pages
VISION_RETRY           = "retry"           # standard first, then vision as second pass

# Annotation-aware extraction prompt — the spec's recommended prompt
VISION_ANNOTATION_PROMPT = (
    "Extract all visible text from these document pages. "
    "Include typed text, handwritten notes, initials, dates, checkmarks, stamps, "
    "margin notes, cross-outs, corrections, highlights, and annotations. "
    "Preserve names, numbers, amounts, and dates exactly as shown. "
    "If handwritten or annotated content appears, clearly identify it as annotation text. "
    "IMPORTANT — epistemic honesty (Item 21): "
    "If text is clearly visible, transcribe it. "
    "If text is partially visible or unclear, prefix it with [UNCLEAR: ...]. "
    "If content is inferred rather than explicitly visible, prefix with [INFERRED: ...]. "
    "Never invent or fill in values you cannot actually see. "
    "Return the text in reading order and separate each page with '--- PAGE BREAK ---'."
)

# Tabular extraction protection note added to all prompts (Item 19)
_TABULAR_GUARD = (
    " If the page contains structured tables with clear columns and rows, "
    "preserve the table structure using tab-separated or pipe-delimited format. "
    "Do not flatten structured tables into plain sentences."
)

# Document-family-specific prompts (Item 8)
VISION_PROMPTS: dict[str, str] = {
    "invoice_receipt": (
        "Extract all text from this invoice/receipt. Capture: vendor name, invoice number, "
        "invoice date, due date, line items with descriptions and amounts, subtotal, tax, "
        "total amount due, payment terms, PO reference number. "
        "Also capture any handwritten approval notes, initials, stamps, or payment confirmations. "
        "Flag any amounts that appear handwritten or corrected."
    ),
    "bank_cash_activity": (
        "Extract all text from this bank statement page. Capture: account holder name, bank name, "
        "account number, statement period (from/to dates), beginning balance, ending balance, "
        "total deposits, total withdrawals, individual transactions with dates/descriptions/amounts. "
        "Also capture any handwritten reconciliation notes, circled items, or margin annotations."
    ),
    "contract_agreement": (
        "Extract all text from this contract or agreement. Capture: party names (lessor/lessee, "
        "grantor/grantee, vendor/client), effective date, expiration/end date, term in months, "
        "payment amounts and schedule, key obligations, schedule or exhibit references. "
        "Also capture handwritten amendments, initials on changes, and any crossed-out text "
        "with the replacement wording."
    ),
    "grant_donor_funding": (
        "Extract all text from this grant award or gift agreement. Capture: grantor/donor name, "
        "recipient organization, award amount, grant number, CFDA number if present, "
        "period of performance (start and end dates), payment schedule, restrictions, "
        "conditions, and reporting requirements. "
        "Also capture any handwritten notes, conditions added by hand, or approval stamps."
    ),
    "governance_approval": (
        "Extract all text from these board minutes or governance document. Capture: meeting date, "
        "attendees, motions with exact wording, vote counts (for/against/abstain), action items, "
        "financial approvals with amounts, related-party disclosures, executive compensation "
        "approvals, and any deferred or tabled items. "
        "Also capture handwritten annotations, added notes, corrections, and initials."
    ),
    "payment_proof": (
        "Extract all text from this payment document, check, or EFT record. Capture: payee name, "
        "payer name, amount, date, check number or reference number, memo line, bank/account "
        "information, and endorsement information if visible. "
        "Also capture any handwritten amounts, dates, approval signatures, or notations."
    ),
    "default": VISION_ANNOTATION_PROMPT,
}


def get_vision_prompt_for_family(document_family: str, has_tables: bool = False) -> str:
    """Return the document-type-specific vision prompt for a given document family.
    Adds tabular structure guard for pages known to have structured data (Item 19).
    """
    base = VISION_PROMPTS.get(document_family, VISION_PROMPTS["default"])
    if has_tables:
        base = base + _TABULAR_GUARD
    return base

VISION_FULL_DOCUMENT_WARN_PAGES = 15  # warn user above this page count


# ── Cache (in-process, keyed by file hash) ────────────────────────────────────
_extraction_cache: dict[str, ParsedDocument] = {}
_ocr_page_cache:  dict[str, str] = {}       # key: file_hash:page_index
_image_page_cache: dict[str, bytes] = {}    # key: file_hash:page_index:dpi


def _file_hash(path: str) -> str:
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


# ── pdfplumber extraction ─────────────────────────────────────────────────────

def _pdfplumber_extract(
    path: str,
    max_pages: int,
    scan_tables: bool,
    table_pages_limit: int = MAX_PDF_PAGES_TABLE_SCAN,
) -> tuple[list[ParsedPage], list[ParsedTable], int]:
    try:
        import pdfplumber
    except ImportError:
        return [], [], 0

    pages: list[ParsedPage] = []
    tables: list[ParsedTable] = []

    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            limit = min(page_count, max_pages)

            for i, page in enumerate(pdf.pages[:limit]):
                text = page.extract_text() or ""
                conf = min(1.0, len(text) / 500) if text else 0.0
                pages.append(ParsedPage(
                    page_number=i + 1,
                    text=text,
                    char_count=len(text),
                    extractor="pdfplumber",
                    confidence=conf,
                ))

                if scan_tables and i < table_pages_limit:
                    for ti, tbl in enumerate(page.extract_tables() or []):
                        if not tbl or len(tbl) < 2:
                            continue
                        headers = tbl[0]
                        valid_h = sum(1 for h in headers if h and str(h).strip())
                        if valid_h >= 2:
                            clean_h = [str(h).strip() if h else f"Col_{j}"
                                       for j, h in enumerate(headers)]
                            rows = [
                                {clean_h[ci] if ci < len(clean_h) else f"Col_{ci}": cell
                                 for ci, cell in enumerate(row)}
                                for row in tbl[1:]
                            ]
                            tables.append(ParsedTable(
                                page_number=i + 1, table_index=ti,
                                headers=clean_h, rows=rows,
                                row_count=len(rows), extractor="pdfplumber",
                            ))
        return pages, tables, page_count
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")
        return [], [], 0


# ── PyPDF2 — shared reader, no reopen per page ────────────────────────────────

def _load_pypdf2_reader(path: str):
    try:
        from PyPDF2 import PdfReader
        with open(path, "rb") as f:
            return PdfReader(io.BytesIO(f.read()))
    except Exception:
        return None


def _pypdf2_page_text(reader, page_index: int) -> str:
    try:
        if reader and page_index < len(reader.pages):
            return reader.pages[page_index].extract_text() or ""
    except Exception:
        pass
    return ""


# ── extractous full-document pass ─────────────────────────────────────────────

def _extractous_full(path: str) -> dict[int, str]:
    try:
        from extractous import Extractor, TesseractOcrConfig
        extractor = (
            Extractor()
            .set_extract_string_max_length(2_000_000)
            .set_ocr_config(TesseractOcrConfig().set_language("eng"))
        )
        text, _ = extractor.extract_file_to_string(str(path))
        if not text or not text.strip():
            return {}
        page_splits = text.split("\x0c")
        return {i + 1: t for i, t in enumerate(page_splits) if t.strip()}
    except Exception as e:
        logger.debug(f"extractous failed: {e}")
        return {}


# ── OCR — shared fitz doc, no reopen per page ────────────────────────────────

def _load_fitz_doc(path: str):
    try:
        import fitz
        return fitz.open(path)
    except Exception:
        return None


def _ocr_page(fitz_doc, page_index: int, dpi: int = 250,
              cache_key_prefix: str = "") -> str:
    """
    OCR a single page. Uses in-process cache keyed by file+page.
    Strategy: psm 6 (uniform block) first — better for tables and forms.
    Falls back to psm 11 (sparse text) if psm 6 yields less content.
    Uses 300 DPI for best table extraction quality.
    """
    ocr_dpi = max(dpi, 300)  # Always use at least 300 DPI for OCR
    ck = f"{cache_key_prefix}:{page_index}:{ocr_dpi}"
    if ck in _ocr_page_cache:
        return _ocr_page_cache[ck]
    try:
        import pytesseract
        from PIL import Image
        page = fitz_doc[page_index]
        pix = page.get_pixmap(dpi=ocr_dpi, alpha=False)
        mode = "RGB" if pix.n < 4 else "RGBA"
        img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)

        # psm 6 = uniform block — better for table/form layouts (invoices, leases, schedules)
        text_psm6 = pytesseract.image_to_string(img, lang="eng", config="--oem 1 --psm 6")
        # psm 11 = sparse text — better for loose/mixed layouts
        text_psm11 = pytesseract.image_to_string(img, lang="eng", config="--oem 1 --psm 11")

        # Pick whichever produced more content
        text = text_psm6 if len(text_psm6.strip()) >= len(text_psm11.strip()) else text_psm11

        if ck:
            _ocr_page_cache[ck] = text
        return text
    except Exception as e:
        logger.debug(f"OCR page {page_index + 1} failed: {e}")
        return ""


# ── Vision — render page images ───────────────────────────────────────────────

def _render_page_images(fitz_doc, page_indices: list[int],
                        dpi: int = 150, file_hash: str | None = None) -> list[bytes]:
    images: list[bytes] = []
    try:
        from PIL import Image
        import io as _io
        for idx in page_indices:
            if idx >= len(fitz_doc):
                continue

            cache_key = f"{file_hash}:{idx}:{dpi}" if file_hash else None
            if cache_key and cache_key in _image_page_cache:
                images.append(_image_page_cache[cache_key])
                continue

            page = fitz_doc[idx]
            pix = page.get_pixmap(dpi=dpi, alpha=False)
            mode = "RGB" if pix.n < 4 else "RGBA"
            img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
            buf = _io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            data = buf.getvalue()
            if cache_key:
                _image_page_cache[cache_key] = data
            images.append(data)
    except Exception as e:
        logger.warning(f"Page render failed: {e}")
    return images


def render_page_image_cached(path: str, page_index: int, dpi: int = 200) -> bytes:
    """
    Render a single PDF page to JPEG bytes.
    Cached by file hash + page index + dpi.
    Used by rescue path in router — independent of shared fitz handle.
    """
    file_hash = _file_hash(path)
    cache_key = f"{file_hash}:{page_index}:{dpi}"
    if cache_key in _image_page_cache:
        return _image_page_cache[cache_key]

    try:
        import fitz
        import io as _io
        from PIL import Image
    except ImportError:
        return b""

    doc = None
    try:
        doc = fitz.open(path)
        page = doc[page_index]
        pix = page.get_pixmap(dpi=dpi, alpha=False)
        mode = "RGB" if pix.n < 4 else "RGBA"
        img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
        buf = _io.BytesIO()
        img.save(buf, format="JPEG", quality=90)
        data = buf.getvalue()
        _image_page_cache[cache_key] = data
        return data
    except Exception:
        return b""
    finally:
        try:
            if doc:
                doc.close()
        except Exception:
            pass


def _vision_weak_pages(
    fitz_doc,
    weak_indices: list[int],
    provider,
    file_hash: str | None = None,
) -> dict[int, str]:
    if not weak_indices or not provider:
        return {}
    if not hasattr(provider, "extract_text_from_page_images"):
        return {}

    cap = min(len(weak_indices), MAX_VISION_PAGES)
    target_indices = weak_indices[:cap]
    images = _render_page_images(fitz_doc, target_indices, file_hash=file_hash)
    if not images:
        return {}

    try:
        prompt = (
            "Extract ALL text from these document pages faithfully. "
            "Preserve all numbers, dates, names, amounts, and terms exactly as written. "
            "Separate pages with '--- PAGE BREAK ---'."
        )
        combined = provider.extract_text_from_page_images(images=images, prompt=prompt)
        page_texts = combined.split("--- PAGE BREAK ---")
        return {
            target_indices[i]: text.strip()
            for i, text in enumerate(page_texts)
            if i < len(target_indices) and text.strip()
        }
    except Exception as e:
        logger.warning(f"Vision extraction failed: {e}")
        return {}


# ── Page improvement scoring and field-aware merge (Items 1, 3, 9) ───────────

import re as _re

_NUM_PATTERN   = _re.compile(r'\b\d[\d,\.]*\b')
_DATE_PATTERN  = _re.compile(
    r'\b(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}|\d{4}[/\-]\d{2}[/\-]\d{2}|'
    r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})\b',
    _re.IGNORECASE
)
_CURRENCY_PATTERN = _re.compile(r'\$[\d,\.]+|\b\d[\d,]*\.\d{2}\b')
_HW_MARKERS    = [
    'initialed', 'approved by', 'signed by', 'noted by', 'reviewed by',
    'n/a', 'ok', 'see attached', 'per', 'ref', 'initials', 'annotated',
    'handwritten', 'written', 'circled', 'checked', 'marked',
]
_ENTITY_HINTS  = [
    'vendor', 'payee', 'grantor', 'donor', 'bank', 'landlord', 'tenant',
    'lessor', 'lessee', 'organization', 'company', 'inc', 'llc', 'corp',
    'foundation', 'trust', 'services', 'associates',
]
_HIGH_RISK_FIELD_PATTERNS = [
    ('amount',   _re.compile(r'\$[\d,\.]+|total|balance|payment|award', _re.I)),
    ('date',     _re.compile(r'\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}|effective|expires|due', _re.I)),
    ('payee',    _re.compile(r'pay(?:able)?\s+to|vendor|grantor|donor|lessor', _re.I)),
    ('term',     _re.compile(r'term|months?|years?|period|duration', _re.I)),
]


def _score_text(text: str) -> dict:
    """Score a text block for information density."""
    if not text:
        return {"chars": 0, "numbers": 0, "dates": 0, "currencies": 0,
                "entities": 0, "hw_markers": 0}
    lower = text.lower()
    return {
        "chars":      len(text),
        "numbers":    len(_NUM_PATTERN.findall(text)),
        "dates":      len(_DATE_PATTERN.findall(text)),
        "currencies": len(_CURRENCY_PATTERN.findall(text)),
        "entities":   sum(1 for h in _ENTITY_HINTS if h in lower),
        "hw_markers": sum(1 for m in _HW_MARKERS if m in lower),
    }


def _compute_page_improvement_score(
    native_text: str,
    vision_text: str,
) -> tuple[float, str, str]:
    """
    Compare native and vision text to determine how much vision improves the page.
    Returns (improvement_score 0-1, merge_decision, reason).
    """
    if not vision_text or not vision_text.strip():
        return 0.0, MERGE_KEEP_NATIVE, ERR_VISION_OUTPUT_EMPTY

    ns = _score_text(native_text or "")
    vs = _score_text(vision_text)

    reasons = []
    score   = 0.0

    # Char delta (raw length improvement)
    char_delta = vs["chars"] - ns["chars"]
    if char_delta > 200:
        score += 0.25
        reasons.append(f"+{char_delta} chars")
    elif char_delta > 30:
        score += 0.10
        reasons.append(f"+{char_delta} chars")

    # Numeric tokens added
    num_delta = vs["numbers"] - ns["numbers"]
    if num_delta > 3:
        score += 0.20
        reasons.append(f"+{num_delta} numeric tokens")
    elif num_delta > 0:
        score += 0.08
        reasons.append(f"+{num_delta} numbers")

    # Dates added
    date_delta = vs["dates"] - ns["dates"]
    if date_delta > 0:
        score += 0.15
        reasons.append(f"+{date_delta} date(s)")

    # Currency/amounts added
    curr_delta = vs["currencies"] - ns["currencies"]
    if curr_delta > 0:
        score += 0.15
        reasons.append(f"+{curr_delta} amount(s)")

    # Entity hints added
    ent_delta = vs["entities"] - ns["entities"]
    if ent_delta > 0:
        score += 0.10
        reasons.append(f"+{ent_delta} entity hint(s)")

    # Handwriting markers present in vision but not native
    hw_delta = vs["hw_markers"] - ns["hw_markers"]
    if hw_delta > 0:
        score += 0.20
        reasons.append(f"handwriting markers: {hw_delta}")

    # Vision is worse
    if vs["chars"] < ns["chars"] * 0.7 and ns["chars"] > 100:
        return 0.0, MERGE_KEEP_NATIVE, "vision shorter than native — keeping native"

    # Determine merge decision
    score = min(1.0, score)
    reason = "; ".join(reasons) if reasons else "no significant improvement"

    if score >= 0.40:
        return score, MERGE_REPLACE, reason
    elif score >= 0.15:
        return score, MERGE_APPEND, reason
    else:
        return score, MERGE_KEEP_NATIVE, ERR_VISION_NO_IMPROVEMENT


def _extract_handwriting_from_vision(vision_text: str) -> tuple[bool, str, float]:
    """
    Scan vision output for handwriting indicators.
    Returns (has_handwriting, handwriting_text, confidence).
    """
    if not vision_text:
        return False, "", 0.0

    lower = vision_text.lower()
    hw_found = any(m in lower for m in _HW_MARKERS)

    # Lines that appear to be handwritten annotations (short, fragmentary)
    hw_lines = []
    for line in vision_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        line_lower = stripped.lower()
        if any(m in line_lower for m in _HW_MARKERS):
            hw_lines.append(stripped)
        elif len(stripped) < 40 and any(c.isupper() for c in stripped) and not stripped.endswith('.'):
            # Short fragmentary line — often a handwritten annotation
            hw_lines.append(stripped)

    handwriting_text = "\n".join(hw_lines[:20])  # cap at 20 lines
    confidence = min(1.0, len(hw_lines) * 0.15) if hw_lines else 0.0

    return (hw_found or bool(hw_lines)), handwriting_text, confidence


def _extract_match_context(text: str, pattern, match_value: str, context_chars: int = 80) -> str:
    """
    Return a short snippet of text surrounding the first occurrence of match_value,
    so reviewers can see the label/field name adjacent to the extracted value.
    Falls back to empty string if not found.
    """
    if not text:
        return ""
    # Search case-insensitively for the match value in the text
    lower_text = text.lower()
    lower_match = match_value.lower()
    idx = lower_text.find(lower_match)
    if idx == -1:
        return ""
    start = max(0, idx - context_chars)
    end   = min(len(text), idx + len(match_value) + context_chars)
    snippet = text[start:end].replace("\n", " ").strip()
    # Prefix ellipsis if we clipped the start
    if start > 0:
        snippet = "…" + snippet
    if end < len(text):
        snippet = snippet + "…"
    return snippet


def _detect_high_risk_field_changes(
    native_text: str,
    vision_text: str,
    page_number: int,
) -> list[dict]:
    """
    Detect if vision changed values for high-risk fields (amounts, dates, payee names).
    Returns list of change records that need reviewer confirmation.
    Each record now includes a context snippet showing the surrounding text so
    reviewers can see what label or field the value belongs to.
    """
    changes = []
    for field_type, pattern in _HIGH_RISK_FIELD_PATTERNS:
        native_matches = set(m.lower() for m in (pattern.findall(native_text or "")))
        vision_matches = set(m.lower() for m in (pattern.findall(vision_text or "")))
        new_values = vision_matches - native_matches
        if new_values:
            top_values = sorted(new_values)[:5]
            # Build a context snippet for each new value from the vision text
            value_contexts = {}
            for val in top_values:
                ctx = _extract_match_context(vision_text or "", pattern, val)
                if ctx:
                    value_contexts[val] = ctx
            changes.append({
                "field_type":        field_type,
                "page":              page_number,
                "new_values":        top_values,
                "value_contexts":    value_contexts,
                "requires_confirmation": True,
            })
    return changes


# ── Page classification (Item 7) ─────────────────────────────────────────────

PAGE_TYPE_COVER          = "cover"
PAGE_TYPE_SIGNATURE      = "signature"
PAGE_TYPE_SUMMARY        = "summary"
PAGE_TYPE_INVOICE_DETAIL = "invoice_detail"
PAGE_TYPE_BANK_TXNS      = "bank_transaction_page"
PAGE_TYPE_AGREEMENT_TERMS= "agreement_terms"
PAGE_TYPE_SCHEDULE       = "schedule"
PAGE_TYPE_HANDWRITTEN    = "handwritten_markup_page"
PAGE_TYPE_RECEIPT        = "receipt_image"
PAGE_TYPE_SUPPORT        = "support_attachment"
PAGE_TYPE_APPENDIX       = "irrelevant_appendix"
PAGE_TYPE_UNKNOWN        = "unknown"

_PAGE_TYPE_SIGNALS = {
    PAGE_TYPE_SIGNATURE:   ["signature", "signed by", "authorized", "witness", "notary", "executed"],
    PAGE_TYPE_INVOICE_DETAIL: ["invoice", "bill to", "item", "quantity", "unit price", "subtotal", "tax"],
    PAGE_TYPE_BANK_TXNS:   ["deposit", "withdrawal", "balance", "transaction", "debit", "credit", "ach", "wire"],
    PAGE_TYPE_AGREEMENT_TERMS: ["whereas", "hereby agrees", "terms and conditions", "article", "section", "covenant"],
    PAGE_TYPE_SCHEDULE:    ["schedule", "exhibit", "attachment", "appendix", "listing", "table of"],
    PAGE_TYPE_RECEIPT:     ["receipt", "thank you", "total paid", "payment received", "card ending"],
    PAGE_TYPE_SUMMARY:     ["executive summary", "overview", "total", "grand total", "year to date"],
    PAGE_TYPE_COVER:       ["cover page", "title page", "table of contents", "prepared by", "confidential"],
}

def classify_page_type(page_text: str) -> str:
    """Classify a page into a type category based on text signals."""
    if not page_text or len(page_text.strip()) < 20:
        return PAGE_TYPE_HANDWRITTEN  # very little text = likely image/handwritten
    lower = page_text.lower()
    scores = {}
    for ptype, signals in _PAGE_TYPE_SIGNALS.items():
        scores[ptype] = sum(1 for s in signals if s in lower)
    best = max(scores, key=scores.get)
    return best if scores[best] >= 2 else PAGE_TYPE_UNKNOWN


# ── Missing-field vision trigger (Item 5) ────────────────────────────────────

# Fields that block canonical quality when missing
_HIGH_VALUE_MISSING_FIELDS = {
    "amounts":     ["total", "balance", "award_amount", "payment", "monthly_charge"],
    "dates":       ["effective_date", "start", "end", "due_date", "period_start"],
    "parties":     ["grantor", "lessor", "vendor", "payee"],
    "identifiers": ["invoice_number", "grant_number", "account_number", "schedule_number"],
}

def identify_pages_needing_vision_for_missing_fields(
    evidence,
    parsed_doc,
) -> tuple[list[int], dict[int, str]]:
    """
    After canonical extraction, identify pages likely to contain missing required fields.
    Returns (page_numbers_to_retry, reasons).
    This implements trigger reason: TRIGGER_MISSING_FIELDS.
    """
    candidates: list[int] = []
    reasons: dict[int, str] = {}

    if not evidence or not parsed_doc:
        return [], {}

    # Determine what's missing
    missing = []
    if not evidence.amounts:
        missing.append("amounts")
    if not evidence.dates:
        missing.append("dates")
    if len(evidence.parties) < 1:
        missing.append("parties")
    if not evidence.identifiers:
        missing.append("identifiers")

    if not missing:
        return [], {}

    # Find pages most likely to contain the missing field types
    for pg in parsed_doc.pages:
        pn = pg.page_number
        text = pg.text or ""
        lower = text.lower()
        page_reasons = []

        for field_type in missing:
            signals = _HIGH_VALUE_MISSING_FIELDS.get(field_type, [])
            if any(s in lower for s in signals):
                page_reasons.append(f"may contain missing {field_type}")
            elif pg.char_count < 300 and field_type in ("amounts", "dates"):
                page_reasons.append(f"low text — {field_type} may be handwritten")

        if page_reasons:
            candidates.append(pn)
            reasons[pn] = f"Page {pn}: " + "; ".join(page_reasons)

    # If nothing targeted, fall back to all low-text pages
    if not candidates:
        for pg in parsed_doc.pages:
            if pg.char_count < 200:
                candidates.append(pg.page_number)
                reasons[pg.page_number] = f"Page {pg.page_number}: low text yield ({pg.char_count} chars)"

    return sorted(set(candidates)), reasons


# ── Retry page identification ─────────────────────────────────────────────────

def identify_pages_needing_vision(
    parsed_doc,
    min_chars_threshold: int = 300,
) -> tuple[list[int], dict[int, str]]:
    """
    Identify pages in a ParsedDocument that are candidates for vision retry.
    Returns (page_numbers, reasons) where reasons maps page_number -> reason string.

    Criteria (any match triggers):
    - extraction failed entirely (0 chars after all extractors)
    - text is below min_chars_threshold (low yield — likely missed content)
    - page used OCR but still below threshold (OCR didn't help enough)
    """
    candidates: list[int] = []
    reasons: dict[int, str] = {}

    if not parsed_doc or not parsed_doc.pages:
        return [], {}

    for pg in parsed_doc.pages:
        pn = pg.page_number
        chars = pg.char_count or 0
        extractor_used = pg.extractor or "none"

        if chars == 0:
            candidates.append(pn)
            reasons[pn] = f"Page {pn}: extraction failed (0 chars)"
        elif chars < min_chars_threshold:
            candidates.append(pn)
            if extractor_used == "ocr":
                reasons[pn] = f"Page {pn}: low text after OCR ({chars} chars) — likely handwriting"
            else:
                reasons[pn] = f"Page {pn}: low text yield ({chars} chars from {extractor_used})"

    return sorted(candidates), reasons


# ── Page selection parser ─────────────────────────────────────────────────────

def parse_page_selection(page_input: str, max_pages: int) -> tuple[list[int], str]:
    """
    Parse a user page selection string into a sorted, deduplicated list of
    1-based page numbers. Returns (pages, error_message).
    Accepts: "3", "1,4,7", "2-5", "1-3,6,8-10"
    Returns 1-based page numbers clamped to [1, max_pages].
    """
    pages: set[int] = set()
    errors: list[str] = []

    if not page_input or not page_input.strip():
        return [], "No pages specified"

    for part in page_input.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            bounds = part.split("-", 1)
            try:
                lo, hi = int(bounds[0].strip()), int(bounds[1].strip())
                if lo < 1 or hi < lo:
                    errors.append(f"Invalid range: {part}")
                    continue
                for p in range(lo, hi + 1):
                    if 1 <= p <= max_pages:
                        pages.add(p)
            except ValueError:
                errors.append(f"Could not parse range: {part}")
        else:
            try:
                p = int(part)
                if 1 <= p <= max_pages:
                    pages.add(p)
                else:
                    errors.append(f"Page {p} out of range (1-{max_pages})")
            except ValueError:
                errors.append(f"Not a number: {part!r}")

    if not pages and not errors:
        return [], "No valid pages found"

    return sorted(pages), "; ".join(errors) if errors else ""


# ── Generalised vision extractor ──────────────────────────────────────────────

def run_vision_on_pages(
    path: str,
    page_numbers: list[int],
    provider,
    prompt: Optional[str] = None,
    dpi: int = 200,
) -> dict:
    """
    Run GPT-4o Vision on an arbitrary set of 1-based page numbers from a PDF.

    Uses structured per-page JSON output for reliable page attribution —
    no dependency on delimiter splitting.

    Returns:
      page_texts:        {page_number: text}          — successfully extracted pages
      page_status:       {page_number: status_str}    — status per requested page
      pages_attempted:   list[int]                    — all pages we tried
      pages_rendered:    list[int]                    — pages where image render succeeded
      pages_succeeded:   list[int]                    — pages with usable text returned
      pages_failed:      list[int]                    — pages with errors or empty output
      prompt_used:       str
      is_partial:        bool                         — True if any requested page failed
      chars_returned:    int
      error:             str | None                   — fatal error (all pages failed)
    """
    import time as _time
    t_start = _time.perf_counter()

    used_prompt = prompt or VISION_ANNOTATION_PROMPT
    result: dict = {
        "page_texts":       {},
        "page_status":      {pn: VISION_PAGE_REQUESTED for pn in page_numbers},
        "pages_attempted":  list(page_numbers),
        "pages_rendered":   [],
        "pages_succeeded":  [],
        "pages_failed":     [],
        "prompt_used":      used_prompt,
        "is_partial":       False,
        "chars_returned":   0,
        "elapsed_seconds":  0.0,
        "error":            None,
    }

    if not page_numbers or not provider:
        result["error"] = "No pages or provider"
        for pn in page_numbers:
            result["page_status"][pn] = VISION_PAGE_FAILED
        return result
    if not hasattr(provider, "extract_text_from_page_images"):
        result["error"] = "Provider does not support vision"
        return result

    file_hash = _file_hash(path)
    zero_indices = [p - 1 for p in page_numbers]

    # For image files, use the raw bytes directly without PDF rendering
    _img_ext = Path(path).suffix.lower()
    _is_image_file = _img_ext in (".jpg", ".jpeg", ".png", ".webp", ".tiff", ".tif", ".bmp")

    if _is_image_file:
        try:
            image_bytes = Path(path).read_bytes()
            images = [image_bytes]  # treat the whole image as "page 1"
            fitz_doc = None
        except Exception as _ie:
            result["error"] = f"Could not read image file: {_ie}"
            for pn in page_numbers:
                result["page_status"][pn] = VISION_PAGE_FAILED
            return result
    else:
        fitz_doc = _load_fitz_doc(path)
        if not fitz_doc:
            result["error"] = "Could not open PDF for rendering"
            for pn in page_numbers:
                result["page_status"][pn] = VISION_PAGE_FAILED
            return result

    try:
        if not _is_image_file:
            images = _render_page_images(fitz_doc, zero_indices, dpi=dpi, file_hash=file_hash)
        if not images:
            result["error"] = "Page rendering produced no images"
            for pn in page_numbers:
                result["page_status"][pn] = VISION_PAGE_FAILED
            return result

        result["pages_rendered"] = list(page_numbers[:len(images)])
        # Mark pages we couldn't render
        for pn in page_numbers[len(images):]:
            result["page_status"][pn] = VISION_PAGE_FAILED
            result["pages_failed"].append(pn)

        # Use structured output prompt — ask for JSON array with explicit page numbers
        # so page attribution does not depend on delimiter splitting
        _pages_str = str(result['pages_rendered'])
        _json_instr = (
            "\n\nReturn your response as a JSON array only, with no other text:\n"
            '[{"page": <1-based page number>, "text": "<extracted text>", '
            '"has_handwriting": <true/false>}, ...]\n'
            f"The pages in this request are (in order): {_pages_str}"
        )
        structured_prompt = used_prompt + _json_instr

        raw = provider.extract_text_from_page_images(
            images=images,
            prompt=structured_prompt,
        )

        # Parse structured JSON response
        parsed_pages = _parse_vision_json_response(raw, page_numbers[:len(images)])

        total_chars = 0
        for pn, text in parsed_pages.items():
            stripped = text.strip()
            if stripped:
                result["page_texts"][pn] = stripped
                result["pages_succeeded"].append(pn)
                result["page_status"][pn] = VISION_PAGE_COMPLETE
                total_chars += len(stripped)
            else:
                result["pages_failed"].append(pn)
                result["page_status"][pn] = VISION_PAGE_FAILED

        # Any rendered page not in parsed_pages also failed
        for pn in result["pages_rendered"]:
            if pn not in result["page_texts"] and pn not in result["pages_failed"]:
                result["pages_failed"].append(pn)
                result["page_status"][pn] = VISION_PAGE_FAILED

        result["chars_returned"] = total_chars
        result["is_partial"]     = len(result["pages_failed"]) > 0

        # Truncation detection: if last rendered page has no output, flag it
        if result["pages_rendered"] and result["pages_rendered"][-1] not in result["page_texts"]:
            result["truncation_suspected"] = True
            logger.warning(
                f"Vision: last page {result['pages_rendered'][-1]} missing from output — "
                f"possible output truncation"
            )

    except Exception as e:
        logger.warning(f"run_vision_on_pages failed: {e}")
        result["error"] = str(e)
        for pn in page_numbers:
            result["page_status"][pn] = VISION_PAGE_FAILED
            if pn not in result["pages_failed"]:
                result["pages_failed"].append(pn)
    finally:
        try:
            if fitz_doc is not None:
                fitz_doc.close()
        except Exception:
            pass
        result["elapsed_seconds"] = round(_time.perf_counter() - t_start, 2)

    return result


def _parse_vision_json_response(raw: str, page_numbers: list[int]) -> dict[int, str]:
    """
    Parse the structured JSON response from the vision model.
    Falls back to delimiter splitting if JSON parsing fails — robust to model variance.
    Returns {page_number: text}.
    """
    import json as _json
    import re as _re

    raw = (raw or "").strip()

    # Attempt 1: parse the whole response as JSON
    try:
        data = _json.loads(raw)
        if isinstance(data, list):
            result = {}
            for item in data:
                if isinstance(item, dict) and "page" in item and "text" in item:
                    pn = int(item["page"])
                    result[pn] = str(item.get("text", ""))
            if result:
                return result
    except (ValueError, TypeError, KeyError):
        pass

    # Attempt 2: extract JSON array from response that has extra prose
    try:
        match = _re.search(r'\[\s*\{.*?\}\s*\]', raw, _re.DOTALL)
        if match:
            data = _json.loads(match.group(0))
            if isinstance(data, list):
                result = {}
                for item in data:
                    if isinstance(item, dict) and "page" in item and "text" in item:
                        pn = int(item["page"])
                        result[pn] = str(item.get("text", ""))
                if result:
                    return result
    except (ValueError, TypeError, KeyError):
        pass

    # Attempt 3: fall back to delimiter splitting with sequential assignment
    # This is the old behavior — used only if JSON parsing failed entirely
    logger.info("Vision: JSON parse failed, falling back to delimiter splitting")
    splits = raw.split("--- PAGE BREAK ---")
    result = {}
    for i, text in enumerate(splits):
        if i < len(page_numbers):
            stripped = text.strip()
            if stripped:
                result[page_numbers[i]] = stripped
    return result


def run_full_document_vision(
    path: str,
    provider,
    prompt: Optional[str] = None,
    dpi: int = 200,
    max_pages: Optional[int] = None,
) -> dict:
    """
    Run vision on every page of a PDF document.
    Processes ALL pages — no silent truncation.
    max_pages: if None, processes the entire document.
                Set to a positive integer to cap (e.g. for testing).
    Returns the same structure as run_vision_on_pages plus:
      - combined_text: all page texts joined in order
      - page_count:    total pages in document
      - pages_requested: how many pages were requested
    """
    try:
        import fitz as _fitz_count
        doc = _fitz_count.open(path)
        total_in_doc = len(doc)
        doc.close()
    except Exception:
        total_in_doc = 0

    total = (min(total_in_doc, max_pages) if max_pages else total_in_doc) or 1
    page_numbers = list(range(1, total + 1))

    logger.info(
        f"run_full_document_vision: {Path(path).name} — "
        f"{total} pages requested (doc has {total_in_doc})"
    )

    result = run_vision_on_pages(path, page_numbers, provider, prompt=prompt, dpi=dpi)
    result["page_count"]       = total_in_doc
    result["pages_requested"]  = len(page_numbers)

    # Build combined text in page order
    texts = []
    for pn in sorted(result["page_texts"].keys()):
        page_text = result["page_texts"][pn]
        texts.append(f"[Page {pn}]\n{page_text}")
    result["combined_text"] = "\n\n".join(texts)

    return result


# ── Non-PDF direct extraction ─────────────────────────────────────────────────

def _extract_direct(path: str) -> ParsedDocument:
    p = Path(path)
    ext = p.suffix.lower()
    tables: list[ParsedTable] = []
    text = ""

    fh = _file_hash(path)
    try:
        if ext in (".csv", ".tsv"):
            import pandas as pd
            df = pd.read_csv(path)
            text = df.to_string(index=False)
            tables = [ParsedTable(
                page_number=1,
                headers=df.columns.tolist(),
                rows=df.head(200).to_dict("records"),
                row_count=len(df),
                extractor="direct",
            )]

        elif ext in (".xlsx", ".xls"):
            import pandas as pd
            df = pd.read_excel(path)
            text = df.to_string(index=False)
            tables = [ParsedTable(
                page_number=1,
                headers=df.columns.tolist(),
                rows=df.head(200).to_dict("records"),
                row_count=len(df),
                extractor="direct",
            )]

        elif ext == ".txt":
            text = p.read_text(encoding="utf-8", errors="replace")

        elif ext == ".docx":
            from docx import Document
            doc = Document(path)
            text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())

        elif ext in (".jpg", ".jpeg", ".png", ".webp", ".tiff", ".tif", ".bmp"):
            # Image files: store bytes and mark as image — vision path handles extraction.
            # _image_bytes persists so run_vision_on_pages can use it directly.
            image_bytes = p.read_bytes()
            pg = ParsedPage(
                page_number=1, text="", char_count=0,
                extractor="direct_image", confidence=0.0, image_used=True,
                native_text="", final_extractor_used="vision",
            )
            return ParsedDocument(
                source_file=p.name,
                file_hash=fh,
                mime_type=f"image/{ext.lstrip('.')}",
                full_text="",
                page_count=1,
                pages=[pg],
                tables=[],
                extraction_chain=["direct_image"],
                primary_extractor="direct_image",
                confidence=0.0,
                warnings=["Image file — vision extraction required"],
                document_specific={"_image_bytes": image_bytes, "_image_ext": ext},
            )

        else:
            return ParsedDocument(
                source_file=p.name,
                errors=[f"Unsupported file type: {ext}"],
            )

        pg = ParsedPage(
            page_number=1, text=text, char_count=len(text),
            extractor="direct", confidence=1.0 if text else 0.0,
        )
        return ParsedDocument(
            source_file=p.name,
            file_hash=fh,
            mime_type=f"text/{ext.lstrip('.')}",
            full_text=text,
            page_count=1,
            pages=[pg],
            tables=[t.model_dump() for t in tables],
            extraction_chain=["direct"],
            primary_extractor="direct",
            confidence=1.0 if text else 0.0,
        )

    except Exception as e:
        return ParsedDocument(source_file=p.name, errors=[str(e)])


# ── Assemble ParsedDocument from pages ───────────────────────────────────────

def _assemble(
    source_file: str,
    file_hash: str,
    pages: list[ParsedPage],
    tables: list[ParsedTable],
    page_count: int,
    extraction_chain: list[str],
    weak_pages: list[int],
    ocr_pages: list[int],
    vision_pages: list[int],
    warnings: list[str],
    errors: list[str],
) -> ParsedDocument:
    full_text = "\n\n".join(
        f"[Page {pg.page_number}]\n{pg.text}"
        for pg in pages if pg.text.strip()
    )
    total_chars = sum(pg.char_count for pg in pages)
    n = max(page_count, len(pages), 1)
    confidence = min(1.0, (total_chars / n) / 500)

    # Compute primary_extractor from actual final page outputs
    extractor_counts: dict[str, int] = {}
    for pg in pages:
        extractor_counts[pg.extractor] = extractor_counts.get(pg.extractor, 0) + pg.char_count
    primary = max(extractor_counts, key=extractor_counts.get) if extractor_counts else "none"

    weak_final = [pg.page_number for pg in pages if pg.char_count < MIN_CHARS_ACCEPTABLE]
    if weak_final:
        warnings.append(f"{len(weak_final)} page(s) below acceptable threshold")

    return ParsedDocument(
        source_file=source_file,
        file_hash=file_hash,
        mime_type="application/pdf",
        full_text=full_text,
        page_count=n,
        pages=pages,
        tables=[t.model_dump() if hasattr(t, "model_dump") else t for t in tables],
        extraction_chain=list(dict.fromkeys(extraction_chain)),
        primary_extractor=primary,
        confidence=confidence,
        weak_pages=weak_final,
        ocr_pages=ocr_pages,
        vision_pages=vision_pages,
        warnings=warnings,
        errors=errors,
    )


# ── Fast lane ─────────────────────────────────────────────────────────────────

def extract_fast(path: str, provider=None, ocr_semaphore=None) -> ParsedDocument:
    """
    Fast Review lane:
    - pdfplumber text (no table scan)
    - shared PyPDF2 for weak pages
    - no OCR, no vision
    - one extractous pass if globally weak
    """
    cache_key = _file_hash(path)
    cache_hit = _extraction_cache.get(f"fast:{cache_key}")
    if cache_hit:
        return cache_hit

    p = Path(path)
    extraction_chain: list[str] = []
    warnings: list[str] = []
    errors: list[str] = []

    # pdfplumber — text only, no table scan
    pages, tables, page_count = _pdfplumber_extract(
        path, MAX_PDF_PAGES_FAST, scan_tables=False
    )
    if pages:
        extraction_chain.append("pdfplumber")

    # Shared PyPDF2 reader for weak pages
    pypdf2_reader = _load_pypdf2_reader(path)
    if pypdf2_reader:
        pypdf2_used = False
        for i, pg in enumerate(pages):
            if pg.char_count < MIN_CHARS_WEAK:
                t = _pypdf2_page_text(pypdf2_reader, i)
                if len(t) > pg.char_count:
                    pages[i] = ParsedPage(
                        page_number=pg.page_number, text=t, char_count=len(t),
                        extractor="pypdf2", confidence=min(1.0, len(t) / 500),
                    )
                    pypdf2_used = True
        if pypdf2_used:
            extraction_chain.append("pypdf2")

    # Global weak check — if still very low overall, try extractous once
    total_chars = sum(pg.char_count for pg in pages)
    n = max(page_count, len(pages), 1)
    if total_chars / n < MIN_CHARS_WEAK:
        ext_pages = _extractous_full(path)
        if ext_pages:
            extraction_chain.append("extractous")
            for pn, t in ext_pages.items():
                idx = pn - 1
                if idx < len(pages) and len(t) > pages[idx].char_count:
                    pages[idx] = ParsedPage(
                        page_number=pn, text=t, char_count=len(t),
                        extractor="extractous", confidence=min(1.0, len(t) / 500),
                    )
            if not pages:
                pages = [
                    ParsedPage(page_number=pn, text=t, char_count=len(t),
                               extractor="extractous", confidence=min(1.0, len(t) / 500))
                    for pn, t in ext_pages.items()
                ]
                page_count = len(pages)

    # Smart escalation — if fast mode is still critically weak, run OCR on worst pages
    total_chars_fast = sum(pg.char_count for pg in pages)
    n_pages = max(page_count, len(pages), 1)
    avg_cpp = total_chars_fast / n_pages
    ocr_pages_fast: list[int] = []
    fitz_doc_fast = None  # initialized here so cleanup pass can safely reference it

    if avg_cpp < MIN_CHARS_CRITICAL and pages:
        logger.info(f"Fast mode critically weak ({avg_cpp:.0f} chars/page) — escalating weak pages to OCR")
        fitz_doc_fast = _load_fitz_doc(path)
        if fitz_doc_fast:
            # Always OCR first 2 pages (most likely to contain key terms/header info)
            # Then fill remaining slots with lowest-char pages
            priority_indices = [i for i in [0, 1] if i < len(pages)
                                and pages[i].char_count < MIN_CHARS_CRITICAL]
            remaining_critical = sorted(
                [i for i, pg in enumerate(pages)
                 if pg.char_count < MIN_CHARS_CRITICAL and i not in priority_indices],
                key=lambda i: pages[i].char_count,
            )
            critical_indices = (priority_indices + remaining_critical)[:MAX_OCR_PAGES]

            ocr_used = False
            fhash = _file_hash(path)
            for i in critical_indices:
                if ocr_semaphore:
                    with ocr_semaphore:
                        t = _ocr_page(fitz_doc_fast, i, cache_key_prefix=fhash)
                else:
                    t = _ocr_page(fitz_doc_fast, i, cache_key_prefix=fhash)
                if len(t.strip()) > pages[i].char_count:
                    pages[i] = ParsedPage(
                        page_number=pages[i].page_number, text=t,
                        char_count=len(t), extractor="ocr",
                        confidence=0.75, image_used=True,
                    )
                    ocr_pages_fast.append(pages[i].page_number)
                    ocr_used = True
            if ocr_used:
                extraction_chain.append("ocr_escalated")

    # Cleanup pass — OCR any pages still critical after the main escalation
    # This catches pages that were bumped by the MAX_OCR_PAGES cap
    still_critical = [
        i for i, pg in enumerate(pages)
        if pg.char_count < MIN_CHARS_CRITICAL and pg.page_number not in ocr_pages_fast
    ]
    if still_critical and fitz_doc_fast:
        fhash2 = _file_hash(path)
        for i in still_critical[:3]:  # Cap cleanup at 3 additional pages
            t = _ocr_page(fitz_doc_fast, i, cache_key_prefix=fhash2)
            if len(t.strip()) > pages[i].char_count:
                pages[i] = ParsedPage(
                    page_number=pages[i].page_number, text=t,
                    char_count=len(t), extractor="ocr",
                    confidence=0.75, image_used=True,
                )
                ocr_pages_fast.append(pages[i].page_number)
        if still_critical:
            extraction_chain.append("ocr_cleanup")

    result = _assemble(
        p.name, cache_key, pages, tables, page_count, extraction_chain,
        [], ocr_pages_fast, [], warnings, errors,
    )
    _extraction_cache[f"fast:{cache_key}"] = result
    return result


# ── Deep lane ─────────────────────────────────────────────────────────────────

def extract_deep(path: str, provider=None, ocr_semaphore=None) -> ParsedDocument:
    """
    Deep Extraction lane:
    - pdfplumber text + targeted table scan
    - shared PyPDF2 reader
    - extractous if globally weak
    - OCR top-N weak pages (shared fitz doc)
    - vision top-N critical pages (shared fitz doc)
    """
    cache_key = _file_hash(path)
    cache_hit = _extraction_cache.get(f"deep:{cache_key}")
    if cache_hit:
        return cache_hit

    p = Path(path)
    extraction_chain: list[str] = []
    warnings: list[str] = []
    errors: list[str] = []
    ocr_pages: list[int] = []
    vision_pages: list[int] = []

    # pdfplumber with table scan
    pages, tables, page_count = _pdfplumber_extract(
        path, MAX_PDF_PAGES_FAST, scan_tables=True,
        table_pages_limit=MAX_PDF_PAGES_TABLE_SCAN,
    )
    if pages:
        extraction_chain.append("pdfplumber")

    # Shared PyPDF2 reader
    pypdf2_reader = _load_pypdf2_reader(path)
    if pypdf2_reader:
        pypdf2_used = False
        for i, pg in enumerate(pages):
            if pg.char_count < MIN_CHARS_WEAK:
                t = _pypdf2_page_text(pypdf2_reader, i)
                if len(t) > pg.char_count:
                    pages[i] = ParsedPage(
                        page_number=pg.page_number, text=t, char_count=len(t),
                        extractor="pypdf2", confidence=min(1.0, len(t) / 500),
                    )
                    pypdf2_used = True
        if pypdf2_used:
            extraction_chain.append("pypdf2")

    # extractous if globally weak
    total_chars = sum(pg.char_count for pg in pages)
    n = max(page_count, len(pages), 1)
    if total_chars / n < MIN_CHARS_WEAK:
        ext_pages = _extractous_full(path)
        if ext_pages:
            extraction_chain.append("extractous")
            for pn, t in ext_pages.items():
                idx = pn - 1
                if idx < len(pages) and len(t) > pages[idx].char_count:
                    pages[idx] = ParsedPage(
                        page_number=pn, text=t, char_count=len(t),
                        extractor="extractous", confidence=min(1.0, len(t) / 500),
                    )
            if not pages:
                pages = [
                    ParsedPage(page_number=pn, text=t, char_count=len(t),
                               extractor="extractous", confidence=min(1.0, len(t) / 500))
                    for pn, t in ext_pages.items()
                ]
                page_count = len(pages)

    # OCR — shared fitz doc, weak pages
    # Always prioritize first 2 pages (most likely to contain key terms/header)
    # then fill remaining slots by lowest char count
    _weak_all = [i for i, pg in enumerate(pages) if pg.char_count < MIN_CHARS_WEAK]
    _priority  = [i for i in [0, 1] if i in _weak_all]
    _remaining = sorted(
        [i for i in _weak_all if i not in _priority],
        key=lambda i: pages[i].char_count,
    )
    weak_indices = (_priority + _remaining)[:MAX_OCR_PAGES]

    if weak_indices:
        fitz_doc = _load_fitz_doc(path)
        if fitz_doc:
            ocr_used = False
            fhash_deep = _file_hash(path)
            for i in weak_indices:
                if ocr_semaphore:
                    with ocr_semaphore:
                        t = _ocr_page(fitz_doc, i, cache_key_prefix=fhash_deep)
                else:
                    t = _ocr_page(fitz_doc, i, cache_key_prefix=fhash_deep)
                if len(t.strip()) > pages[i].char_count:
                    pages[i] = ParsedPage(
                        page_number=pages[i].page_number, text=t,
                        char_count=len(t), extractor="ocr",
                        confidence=0.75, image_used=True,
                    )
                    ocr_pages.append(pages[i].page_number)
                    ocr_used = True

            if ocr_used:
                extraction_chain.append("ocr")

            # Cleanup pass — OCR any pages still critical after the main pass
            still_critical = [
                i for i, pg in enumerate(pages)
                if pg.char_count < MIN_CHARS_CRITICAL
                and pg.page_number not in ocr_pages
            ]
            for i in still_critical[:3]:
                if ocr_semaphore:
                    with ocr_semaphore:
                        t = _ocr_page(fitz_doc, i, cache_key_prefix=fhash_deep)
                else:
                    t = _ocr_page(fitz_doc, i, cache_key_prefix=fhash_deep)
                if len(t.strip()) > pages[i].char_count:
                    pages[i] = ParsedPage(
                        page_number=pages[i].page_number, text=t,
                        char_count=len(t), extractor="ocr",
                        confidence=0.75, image_used=True,
                    )
                    ocr_pages.append(pages[i].page_number)
            if still_critical:
                extraction_chain.append("ocr_cleanup")

            # Vision — still critical after OCR, top pages only
            critical_indices = sorted(
                [i for i, pg in enumerate(pages) if pg.char_count < MIN_CHARS_CRITICAL],
                key=lambda i: pages[i].char_count,
            )[:MAX_VISION_PAGES]

            if critical_indices and provider is not None:
                vision_results = _vision_weak_pages(fitz_doc, critical_indices, provider, file_hash=fhash_deep)
                if vision_results:
                    extraction_chain.append("vision")
                    for i, t in vision_results.items():
                        if i < len(pages) and len(t) > pages[i].char_count:
                            pages[i] = ParsedPage(
                                page_number=pages[i].page_number, text=t,
                                char_count=len(t), extractor="vision",
                                confidence=0.85, image_used=True,
                            )
                            vision_pages.append(pages[i].page_number)

    result = _assemble(
        p.name, cache_key, pages, tables, page_count, extraction_chain,
        [], ocr_pages, vision_pages, warnings, errors,
    )
    _extraction_cache[f"deep:{cache_key}"] = result
    return result


# ── Main entry point ──────────────────────────────────────────────────────────

def extract(
    path: str,
    provider=None,
    mode: str = "fast",
    ocr_semaphore=None,
) -> ParsedDocument:
    """
    Main extraction entry point.
    mode="fast"  → extract_fast (no OCR/vision unless critically weak)
    mode="deep"  → extract_deep (OCR + vision on weak pages)
    Non-PDF files always use direct extraction.
    ocr_semaphore: threading.Semaphore passed from router for OCR throttling.
    """
    p = Path(path)
    if p.suffix.lower() not in (".pdf",):
        return _extract_direct(path)

    if mode == "deep":
        return extract_deep(path, provider=provider, ocr_semaphore=ocr_semaphore)
    return extract_fast(path, provider=provider, ocr_semaphore=ocr_semaphore)
